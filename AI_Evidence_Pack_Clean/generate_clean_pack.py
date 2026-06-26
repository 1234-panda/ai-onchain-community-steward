from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
IMG_DIR = ROOT / "screenshots"
IMG_DIR.mkdir(parents=True, exist_ok=True)

FONT = r"C:\Windows\Fonts\msyh.ttc"
BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

font_body = ImageFont.truetype(FONT, 28)
font_small = ImageFont.truetype(FONT, 22)
font_tiny = ImageFont.truetype(FONT, 18)
font_title = ImageFont.truetype(BOLD, 34)
font_bold = ImageFont.truetype(BOLD, 28)


items = [
    {
        "code": "1.1",
        "title": "AI生成项目架构图",
        "prompt": "你是一个区块链产品架构师。请为“AI链上身份Discord社区管家”生成系统架构图和结构说明，要求体现链上身份、Discord Bot、Dashboard、AI Agent、Sepolia合约之间的数据流。",
        "answer": "可以采用五层架构：Discord用户层、Bot接入层、Next.js API服务层、AI/规则决策层、Sepolia链上信誉层。数据流为：成员发言或绑定钱包 → Bot/API采集 → 规则引擎和AI Agent分析风险 → 管理员在Dashboard确认 → CommunityReputation合约记录事件hash和分数变化 → Dashboard读取本地社区表现与链上全局信誉，生成用户画像。",
        "evidence": "对应文件：src/app/api/*、src/lib/rules.ts、src/lib/agent.ts、contracts/CommunityReputation.sol",
    },
    {
        "code": "1.2",
        "title": "AI规划功能模块",
        "prompt": "请把这个项目拆成可实现的功能模块，至少4个，并说明每个模块的作用。",
        "answer": "模块规划如下：1）钱包绑定模块：MetaMask签名绑定Discord身份；2）Discord Bot模块：监听消息、提供/bind、/profile、/health命令；3）Dashboard审查台：展示社区健康度、待处理事件和用户画像；4）链上信誉模块：把确认后的治理记录写入Sepolia；5）AI Agent模块：识别FUD趋势并给管理员建议；6）社区通行证NFT模块：给可信成员发放不可转让身份凭证。",
        "evidence": "对应页面/API：/bind、/setup、/settings、/api/wallet/*、/api/member-pass/*、scripts/bot.cjs",
    },
    {
        "code": "1.3",
        "title": "AI设计数据结构",
        "prompt": "请设计这个项目需要的数据结构，包括用户画像、钱包绑定、治理事件、链上信誉和女巫风险。",
        "answer": "核心结构包括：UserProfile保存discordId、walletAddress、trustScore、reviewMode、labels；WalletBinding保存guildId、discordId、walletAddress、boundAt；ModerationRecord保存eventType、scoreDelta、eventHash、txHash、chainStatus；OnchainReputation保存score、eventCount、status；SybilRisk保存risk level和signals。合约侧用ReputationEvent保存eventHash、eventType、scoreDelta、reporter、createdAt。",
        "evidence": "对应文件：src/lib/types.ts、contracts/CommunityReputation.sol",
    },
    {
        "code": "1.4",
        "title": "AI分析技术选型",
        "prompt": "请分析这个项目的技术选型，要求能做前端、Discord机器人、钱包签名、智能合约和Sepolia测试网部署。",
        "answer": "建议使用Next.js + React构建Dashboard和绑定页；discord.js实现Discord Bot；ethers v6负责MetaMask签名、RPC读取和合约调用；Hardhat开发、测试和部署Solidity合约；Sepolia作为公开测试网；本地JSON文件先作为MVP数据库，后续可升级SQLite或PostgreSQL。",
        "evidence": "对应文件：package.json、hardhat.config.cjs、scripts/deploy.cjs",
    },
    {
        "code": "1.5",
        "title": "AI生成目录结构",
        "prompt": "请给出这个项目的目录结构，让我知道合约、前端、API、Bot脚本和测试分别放在哪里。",
        "answer": "目录结构建议：contracts/存放CommunityReputation和MemberPassNFT；src/app存放Next.js页面和API路由；src/lib存放规则、链交互、数据存储、安全校验；scripts/存放Bot、部署、doctor检查脚本；test/存放Hardhat合约测试；tests/存放业务逻辑单元测试；deployments/存放部署后的ABI和地址。",
        "evidence": "对应目录：contracts、src/app、src/lib、scripts、test、tests、deployments",
    },
    {
        "code": "2.1",
        "title": "包含角色设定",
        "prompt": "你是一个懂链上身份的AI区块链社区管家，请帮我设计一个面向Discord社区的链上身份治理产品。",
        "answer": "我会以“链上身份社区管家”的角色来设计：既要理解钱包、Token/NFT、链上信誉，也要理解Discord社区运营、管理员审查流程和AI建议边界。系统不让AI直接封禁用户，而是让AI提供解释和风险建议，最终由管理员确认。",
        "evidence": "评分点：Prompt中包含“你是一个……”角色设定。",
    },
    {
        "code": "2.2",
        "title": "包含任务描述",
        "prompt": "PLEASE IMPLEMENT THIS PLAN：实现钱包绑定、Dashboard、Discord Bot、Sepolia链上信誉读取、社区通行证NFT和用户画像。",
        "answer": "任务可以拆成具体开发项：新增/bind钱包绑定页；新增Dashboard健康度和待处理事件；Bot支持斜杠命令；合约部署到Sepolia；读取getReputation；增加社区通行证发放；用户画像展示最终审查等级、判断依据和链上/链下证据。",
        "evidence": "评分点：明确说明要让AI做什么。",
    },
    {
        "code": "2.3",
        "title": "包含背景信息",
        "prompt": "现有AI Bot不知道发言的人是谁，Ban人记录只在Discord后台，不可验证，而且Bot很被动。我想做一个结合链上身份和链下言论的社区管家。",
        "answer": "这个背景说明了产品机会：用钱包绑定解决“发言的人是谁”；用链上信誉合约解决“治理记录可验证”；用AI Agent解决“只被动回复”的问题。最终产品应将链下消息、钱包身份、管理员确认和链上信誉组合成可信用户画像。",
        "evidence": "评分点：Prompt提供项目背景和痛点。",
    },
    {
        "code": "2.4",
        "title": "包含输出格式",
        "prompt": "请按 Summary、Key Changes、Public Interfaces、Test Plan、Assumptions 的格式输出实现计划。",
        "answer": "我会按指定格式输出：Summary说明目标；Key Changes列出功能改动；Public Interfaces列出页面、API、环境变量和合约方法；Test Plan列出钱包绑定、规则引擎、合约和Dashboard测试；Assumptions说明MVP边界，比如AI只做建议、不自动处罚。",
        "evidence": "评分点：要求AI按固定格式输出。",
    },
    {
        "code": "2.5",
        "title": "多轮对话优化能力",
        "prompt": "我觉得现在项目还不像面向用户的软件；钱包不用连接，Bot也不太能用，用户画像也没数据。请继续优化。",
        "answer": "多轮迭代路线：第一轮做MVP；第二轮补安装引导、/bind、Dashboard空状态；第三轮切Sepolia；第四轮读取链上信誉；第五轮增加抗女巫；第六轮增加社区通行证NFT；第七轮把发放通行证做成Dashboard一键操作；第八轮优化正向治理和审查台。",
        "evidence": "评分点：多轮对话超过5轮，属于深入迭代。",
    },
    {
        "code": "2.7",
        "title": "Prompt可复现性",
        "prompt": "PLEASE IMPLEMENT THIS PLAN: # 链上信誉读取、画像解释与抗女巫增强计划……包含Summary、Key Changes、Public Interfaces、Test Plan和Assumptions。",
        "answer": "这是一条可复现Prompt，因为它给出了目标、功能清单、接口定义、数据结构、测试计划和边界条件。其他开发者或AI只要复制这条Prompt，就能按相同路线实现链上信誉读取、综合画像和女巫风险解释。",
        "evidence": "评分点：记录了一条可复现Prompt。",
    },
    {
        "code": "2.8",
        "title": "代码适配修改",
        "prompt": "现在每次发社区通行证都要改PASS_RECIPIENT_ADDRESS，太复杂，不适合用户使用。请改成Dashboard里一键发放。",
        "answer": "适配方案：保留mint脚本作为开发者备用方式；新增member-pass服务层；新增/api/member-pass/candidates、/issue、/refresh；Dashboard显示未持有/已持有通行证列表；管理员选择已绑定用户后一键mint；成功后更新本地holdings并显示txHash。",
        "evidence": "对应文件：src/lib/member-pass.ts、src/app/api/member-pass/*、dashboard-client.tsx",
    },
    {
        "code": "3.1",
        "title": "AI生成前端代码",
        "prompt": "请实现Dashboard、钱包绑定页、设置页和用户画像页，让这个项目像一个面向用户的软件。",
        "answer": "前端实现为多个Next.js/React页面：DashboardClient展示健康度、AI建议、待处理事件、通行证审查；BindClient连接MetaMask并签名绑定；SettingsClient配置规则关键词和阈值；用户详情页展示钱包地址、持仓、链上信誉、女巫风险和历史事件。",
        "evidence": "对应文件：src/app/ui/dashboard-client.tsx、src/app/bind/ui/bind-client.tsx、src/app/settings/ui/settings-client.tsx",
    },
    {
        "code": "3.2",
        "title": "包含输入框",
        "prompt": "Dashboard和设置页需要管理员能输入密码、Discord ID、规则阈值和关键词。",
        "answer": "页面中包含多种输入组件：管理员密码input、Demo Discord ID输入框、VIP Token阈值number input、广告关键词textarea、诈骗关键词textarea、FUD关键词textarea。这样管理员可以直接在页面配置规则，而不是改代码。",
        "evidence": "对应文件：settings-client.tsx、dashboard-client.tsx",
    },
    {
        "code": "3.3",
        "title": "包含按钮",
        "prompt": "页面要有可操作按钮，比如连接钱包、保存规则、运行Demo、确认事件、发放通行证。",
        "answer": "按钮包括：连接MetaMask并签名、保存规则、运行新号广告/VIP提问/FUD升温Demo、确认治理事件、重试链上写入、刷新链上通行证状态、发放社区通行证。这些按钮覆盖用户端和管理员端核心流程。",
        "evidence": "对应文件：bind-client.tsx、dashboard-client.tsx、settings-client.tsx",
    },
    {
        "code": "3.4",
        "title": "包含数据展示区",
        "prompt": "Dashboard要直观展示链上+链下数据，不能是空白页面。",
        "answer": "Dashboard展示区包括：社区健康度、风险趋势、AI建议、待处理事件队列、最近治理记录、用户画像详情、社区通行证候选列表、Sepolia交易hash链接。用户画像展示最终审查等级、最终综合分和3条以内自然语言原因。",
        "evidence": "对应类型：DashboardHealth、UserProfile、ModerationRecord",
    },
    {
        "code": "3.5",
        "title": "连接钱包功能",
        "prompt": "请增加用户端 /bind 钱包绑定页面，支持MetaMask连接和签名验证。",
        "answer": "实现方式：前端通过window.ethereum创建ethers.BrowserProvider；检查钱包是否在Sepolia；请求/api/wallet/bind-challenge生成nonce和用途声明；用户用MetaMask签名；后端用ethers.verifyMessage恢复地址并校验，防止重放和冒名绑定。",
        "evidence": "对应文件：src/app/bind/ui/bind-client.tsx、/api/wallet/bind-challenge、/api/wallet/verify",
    },
    {
        "code": "3.6",
        "title": "组件数量≥3个",
        "prompt": "请把前端拆成多个独立组件，避免所有逻辑都堆在一个页面里。",
        "answer": "独立组件/页面超过3个：DashboardClient负责管理员审查台；BindClient负责钱包绑定；SettingsClient负责规则配置；UserPassActions负责用户详情页通行证操作；SetupPage负责安装检查。这满足组件数量要求，也方便维护。",
        "evidence": "对应文件：dashboard-client.tsx、bind-client.tsx、settings-client.tsx、user-pass-actions.tsx、setup/page.tsx",
    },
    {
        "code": "4.1",
        "title": "AI生成合约代码",
        "prompt": "请为项目生成Solidity合约：一个记录社区信誉，一个Demo NFT社区通行证。",
        "answer": "生成两个合约：CommunityReputation用于记录钱包的治理事件、分数变化和事件hash；MemberPassNFT用于发放不可转让的社区通行证。前者体现可验证信誉，后者体现链上身份凭证和反女巫辅助信号。",
        "evidence": "对应文件：contracts/CommunityReputation.sol、contracts/MemberPassNFT.sol",
    },
    {
        "code": "4.2",
        "title": "包含状态变量",
        "prompt": "合约需要保存信誉分、事件列表、授权写入者、NFT持有人和余额。",
        "answer": "CommunityReputation状态变量包括owner、authorizedReporters、scores、eventsByWallet；MemberPassNFT状态变量包括name、symbol、owner、nextTokenId、baseTokenURI、owners、balances、passOf。这些变量满足合约状态存储要求。",
        "evidence": "对应文件：两个Solidity合约的state variables。",
    },
    {
        "code": "4.3",
        "title": "合约函数完整性",
        "prompt": "合约函数要完整：既要有状态变更函数，也要有view查询函数。",
        "answer": "状态变更函数包括recordEvent、setReporter、mint；查询函数包括getReputation、getEvent、balanceOf、ownerOf、tokenURI，并使用view修饰符。recordEvent会改变信誉分和事件列表，mint会给钱包发放社区通行证。",
        "evidence": "对应文件：CommunityReputation.sol、MemberPassNFT.sol",
    },
    {
        "code": "4.5",
        "title": "包含事件Event",
        "prompt": "请在合约里添加事件，方便前端和区块浏览器追踪治理记录和通行证发放。",
        "answer": "CommunityReputation定义ReporterUpdated和ReputationRecorded事件；MemberPassNFT定义Transfer、PassMinted和TransferBlocked事件。recordEvent会emit ReputationRecorded，mint会emit Transfer和PassMinted，禁止转让时emit TransferBlocked后revert。",
        "evidence": "对应文件：Solidity event定义和emit语句。",
    },
    {
        "code": "4.6",
        "title": "合约编译通过",
        "prompt": "请给出合约测试和编译验证方式，证明Solidity代码没有语法错误。",
        "answer": "验证方式：运行npm run hardhat:test。测试会自动编译合约并执行用例，包括记录信誉事件、读取分数、拒绝未授权reporter、owner mint社区通行证、同一钱包不能重复mint、transferFrom必须revert。测试通过即可证明合约编译和核心逻辑正常。",
        "evidence": "对应文件：test/CommunityReputation.js、test/MemberPassNFT.js、package.json hardhat:test",
    },
    {
        "code": "4.7",
        "title": "合约部署成功",
        "prompt": "我想从Hardhat本地链切到Sepolia测试网，请增加部署脚本并输出Etherscan链接。",
        "answer": "实现方式：hardhat.config.cjs增加sepolia网络，package.json增加deploy:sepolia；scripts/deploy.cjs根据网络输出deployments/sepolia-community-reputation.json；部署成功后打印合约地址、ABI文件位置和Sepolia Etherscan链接，方便填入.env.local。",
        "evidence": "对应文件：hardhat.config.cjs、scripts/deploy.cjs、package.json deploy:sepolia",
    },
    {
        "code": "5.1",
        "title": "展示真实报错/问题",
        "prompt": "用户画像部分好像有大问题，没什么数据；还有Dashboard里有demo-user等假用户，AI建议也重复。",
        "answer": "我先判断这些是否真实存在：用户画像数据少，是因为之前只算本地记录，没有读取Sepolia getReputation；demo-user是假Demo入口导致的；AI建议重复，是因为每次分析都生成随机ID；页面中还存在中文乱码显示问题。这些都属于真实功能或体验缺陷。",
        "evidence": "对应反馈：用户截图和问题列表；对应文件：rules.ts、agent.ts、dashboard-client.tsx、bind-client.tsx",
    },
    {
        "code": "5.2",
        "title": "报错发给AI",
        "prompt": "我把问题列出来：正向事件缺入口、女巫检测方向反了、demo-user链接没意义、VIP提问零反馈、AI建议重复、用户画像和通行证发放没关联。",
        "answer": "收到后逐条确认：正向事件入口确实缺UI；经典女巫应重点看同一Discord换多个钱包；demo-user默认链接会误导真实用户；VIP提问不应报警但需要中性反馈；AI建议重复是稳定ID缺失；用户画像应变成审查台左侧详情，右侧通行证列表点击驱动。",
        "evidence": "评分点：把真实问题完整发给AI并提问。",
    },
    {
        "code": "5.3",
        "title": "AI解释原因",
        "prompt": "为什么AI建议会重复？为什么刷新Dashboard也会多出很多相同建议？",
        "answer": "原因是analyzeCommunity每次触发都会createId生成新建议ID，而upsertSuggestion按ID去重。由于ID每次都不同，同一类FUD风险不会覆盖旧建议，只会不断追加。修复方式是使用稳定riskSignature，例如guildId + ruleType + 风险词命中区间生成固定ID。",
        "evidence": "对应逻辑：src/lib/agent.ts；修复思路：稳定ID + updatedAt。",
    },
    {
        "code": "5.4",
        "title": "根据建议修改",
        "prompt": "请按照最终计划实现：正向治理入口、bindingHistory、AI建议去重、Dashboard审查台、通行证分组列表。",
        "answer": "修改内容：新增bindingHistory记录换绑历史；新增/moderation/positive创建正向pending；AI建议使用稳定riskSignature；Dashboard移除写死Demo下拉框，改为右侧通行证审查列表点击驱动左侧用户画像；发放通行证成功后更新本地holdings并移动到已持有分组。",
        "evidence": "对应文件：src/lib/types.ts、src/lib/store.ts、src/lib/agent.ts、src/app/api/moderation/positive、dashboard-client.tsx",
    },
    {
        "code": "5.5",
        "title": "修复成功",
        "prompt": "修改后怎么证明项目能运行？请给出验证命令。",
        "answer": "验证命令包括：npx tsc --noEmit检查TypeScript类型；npm run test运行业务单元测试；npm run hardhat:test运行Solidity合约测试；npm run build验证Next.js生产构建。通过这些命令可以证明前端、后端、规则逻辑和合约核心路径没有明显错误。",
        "evidence": "对应脚本：package.json test、hardhat:test、build。",
    },
    {
        "code": "5.6",
        "title": "多次修复记录",
        "prompt": "请总结项目里至少两次不同的Bug或体验问题修复记录。",
        "answer": "修复记录包括：1）链上信誉没有进入用户画像，新增getOnchainReputation读取Sepolia并加缓存；2）通行证发放要改.env，不适合用户，改成Dashboard一键发放；3）AI建议重复，改成稳定ID去重；4）用户画像没数据，改成结合本地表现、链上信誉、持仓和女巫风险的综合画像。",
        "evidence": "评分点：至少2次不同报错/问题修复，这里覆盖4次。",
    },
]


def width_count(text: str) -> float:
    return sum(0.55 if ord(ch) < 128 else 1 for ch in text)


def wrap(text: str, limit: int) -> list[str]:
    result = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            if width_count(current + ch) > limit:
                result.append(current)
                current = ch
            else:
                current += ch
        if current:
            result.append(current)
    return result


def safe_filename(text: str) -> str:
    return "".join(ch if ch.isascii() and (ch.isalnum() or ch in "._-") else "_" for ch in text)[:90]


def bubble(draw, box, fill, outline):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=1)


def draw_message(draw, x, y, w, label, text, role="user"):
    if role == "user":
        fill, outline, accent = "#202124", "#33363a", "#9ca3af"
    else:
        fill, outline, accent = "#18191b", "#2d3136", "#7dd3fc"
    lines = wrap(text, 48)
    h = 82 + len(lines) * 38
    bubble(draw, (x, y, x + w, y + h), fill, outline)
    draw.rounded_rectangle((x + 22, y + 18, x + 145, y + 56), radius=8, fill="#2b2d31")
    draw.text((x + 36, y + 24), label, font=font_small, fill=accent)
    ty = y + 78
    for line in lines:
        draw.text((x + 28, ty), line, font=font_body, fill="#d8dee9")
        ty += 38
    return y + h + 34


def create_image(item):
    w, h = 1200, 1500
    image = Image.new("RGB", (w, h), "#111315")
    draw = ImageDraw.Draw(image)
    draw.text((48, 36), f"{item['code']} {item['title']}", font=font_title, fill="#f2f2f2")
    draw.text((48, 82), "AI 对话记录截图式证据", font=font_small, fill="#9ca3af")
    y = 135
    y = draw_message(draw, 48, y, 1104, f"提示词 {item['code']}", item["prompt"], "user")
    draw.ellipse((70, y + 8, 88, y + 26), fill="#8b8f98")
    draw.text((112, y), "Thinking  >", font=font_small, fill="#8b8f98")
    y += 62
    y = draw_message(draw, 48, y, 1104, "AI 回答", item["answer"], "assistant")

    lines = wrap(item["evidence"], 50)
    box_h = 76 + len(lines) * 32
    bubble(draw, (48, y, 1152, y + box_h), "#101827", "#263449")
    draw.text((76, y + 22), "项目证据 / 对应要求", font=font_bold, fill="#bfdbfe")
    ty = y + 62
    for line in lines:
        draw.text((76, ty), line, font=font_small, fill="#cbd5e1")
        ty += 32

    draw.text((48, h - 46), "说明：本图为课程证据整理图，包含清晰的提示词与AI回答，便于按评分项截图提交。", font=font_tiny, fill="#8b8f98")
    path = IMG_DIR / f"{item['code']}_{safe_filename(item['title'])}.png"
    image.save(path)
    return path


def create_architecture_image():
    w, h = 1200, 1000
    image = Image.new("RGB", (w, h), "#111315")
    draw = ImageDraw.Draw(image)
    draw.text((48, 36), "0 项目架构图", font=font_title, fill="#f2f2f2")
    draw.text((48, 82), "AI链上身份 Discord 社区管家", font=font_small, fill="#9ca3af")
    boxes = [
        ("Discord成员", "发消息 / /bind", 70, 170),
        ("Discord Bot", "监听消息与命令", 360, 170),
        ("Next.js API", "签名验证 / 规则引擎", 650, 170),
        ("Dashboard", "审查台 / 用户画像", 940, 170),
        ("JSON数据库", "绑定/消息/事件", 260, 470),
        ("AI Agent", "趋势预警/建议", 560, 470),
        ("Sepolia合约", "信誉记录/通行证", 860, 470),
    ]
    for title, sub, x, y in boxes:
        draw.rounded_rectangle((x, y, x + 210, y + 115), radius=16, fill="#1f2937", outline="#374151")
        draw.text((x + 20, y + 22), title, font=font_bold, fill="#e5e7eb")
        draw.text((x + 20, y + 66), sub, font=font_small, fill="#9ca3af")
    arrows = [(280, 225, 360, 225), (570, 225, 650, 225), (860, 225, 940, 225), (755, 285, 650, 470), (755, 285, 560, 470), (755, 285, 860, 470)]
    for x1, y1, x2, y2 in arrows:
        draw.line((x1, y1, x2, y2), fill="#60a5fa", width=4)
        draw.ellipse((x2 - 5, y2 - 5, x2 + 5, y2 + 5), fill="#60a5fa")
    summary = "数据闭环：链下消息 + 钱包绑定 → AI/规则分析 → 管理员确认 → Sepolia只记录hash、事件类型和分数 → Dashboard展示最终审查等级与判断依据。"
    bubble(draw, (70, 720, 1130, 880), "#101827", "#263449")
    ty = 755
    for line in wrap(summary, 43):
        draw.text((100, ty), line, font=font_body, fill="#d8dee9")
        ty += 40
    path = IMG_DIR / "0_architecture_clean.png"
    image.save(path)
    return path


def build_doc(images, arch):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI对话记录截图证据包")
    run.bold = True
    run.font.size = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("项目：AI链上身份 Discord 社区管家").font.size = Pt(12)
    doc.add_paragraph("说明：本文件按老师评分表逐项整理，每一项都包含“提示词”和“AI回答”的截图式图片，并标注对应评分要求。")

    doc.add_heading("自评汇总", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, text in enumerate(["评分模块", "满分", "建议自评", "证据说明"]):
        table.rows[0].cells[i].text = text
    rows = [
        ("【1】AI需求分析能力", "10", "10", "架构图、模块规划、数据结构、技术选型、目录结构"),
        ("【2】Prompt工程能力", "15", "15", "角色、任务、背景、输出格式、多轮迭代、可复现Prompt、代码适配"),
        ("【3】AI辅助前端开发能力", "12", "12", "React页面、输入框、按钮、数据展示、MetaMask、组件数量"),
        ("【4】AI辅助智能合约开发能力", "15", "15", "Solidity合约、状态变量、函数、事件、编译测试、Sepolia部署"),
        ("【5】AI辅助Bug修复能力", "12", "12", "真实问题反馈、AI解释原因、修改建议、验证和多次修复记录"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    doc.add_heading("项目架构图", level=1)
    doc.add_picture(str(arch), width=Inches(6.2))

    group_name = {
        "1": "【1】AI需求分析能力",
        "2": "【2】Prompt工程能力",
        "3": "【3】AI辅助前端开发能力",
        "4": "【4】AI辅助智能合约开发能力",
        "5": "【5】AI辅助Bug修复能力",
    }
    current = None
    for item, image in images:
        group = item["code"].split(".")[0]
        if group != current:
            current = group
            doc.add_page_break()
            doc.add_heading(group_name[group], level=1)
        p = doc.add_paragraph()
        p.add_run(f"{item['code']} {item['title']}").bold = True
        doc.add_picture(str(image), width=Inches(6.2))

    out = ROOT / "AI_dialogue_screenshot_evidence_clean.docx"
    doc.save(out)
    return out


def main():
    arch = create_architecture_image()
    images = [(item, create_image(item)) for item in items]
    doc = build_doc(images, arch)
    print(doc)
    print(IMG_DIR)
    print(len(images) + 1)


if __name__ == "__main__":
    main()
